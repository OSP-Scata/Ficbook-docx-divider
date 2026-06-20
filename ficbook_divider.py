import os
import tkinter as tk
from tkinter import filedialog
import docx


def open_file():  # открытие файла
    ready['text'] = ''
    selected_file['text'] = ''
    file = filedialog.askopenfile(
        mode='r', filetypes=[('Документ Word', '*.docx')])
    if file:
        global filepath
        global name
        global directory
        filepath = os.path.abspath(file.name)
        directory = os.path.dirname(filepath)
        filename = os.path.basename(file.name)
        name, ext = os.path.splitext(os.path.basename(filepath))
        selected_file['text'] = filename


def submit():  # функция, разделяющая файл
    divider_text = divider.get().rstrip('\n')
    if not divider_text:
        ready['text'] = 'Введите разделитель!'
    else:
        try:
            document = docx.Document(filepath)
            new_doc = docx.Document()
            style = new_doc.styles['Normal']
            style.paragraph_format.left_indent = docx.shared.Cm(-0.5)
            style.paragraph_format.first_line_indent = docx.shared.Cm(0.5)
            part_num = 1
            filename_folder = os.path.join(directory, name)
            os.makedirs(filename_folder, exist_ok=True)

            # сохранение исходного форматирования
            for elem in document.element.body:
                if isinstance(elem, docx.oxml.text.paragraph.CT_P):
                    para = docx.text.paragraph.Paragraph(elem, document)
                    if divider_text in para.text:
                        new_doc.save(f"{filename_folder}/part_{part_num}.docx")
                        new_doc = docx.Document()
                        style = new_doc.styles['Normal']
                        style.paragraph_format.left_indent = docx.shared.Cm(
                            -0.5)
                        style.paragraph_format.first_line_indent = \
                            docx.shared.Cm(0.5)
                        part_num += 1
                        continue
                    new_para = new_doc.add_paragraph()
                    for run in para.runs:
                        new_run = new_para.add_run(run.text)
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.font.size = run.font.size
                        new_run.font.name = run.font.name

                # если в тексте попалась таблица
                elif isinstance(elem, docx.oxml.table.CT_Tbl):
                    pass

            # сохранение файлов
            new_doc.save(f"{filename_folder}/part_{part_num}.docx")
            ready['text'] = 'Готово!'
        except Exception as e:
            print(e)


# GUI
root = tk.Tk()
root.title('Разделение DOCX-файла на главы')
root.geometry('400x275')

import_button = tk.Button(root, text='Импорт .docx', command=open_file)
import_button.pack(padx=6, pady=6)

selected = tk.Label(text='Выбранный файл:')
selected.pack(padx=3, pady=3)
selected_file = tk.Label(root)
selected_file.pack(padx=6, pady=6)

label = tk.Label(root, text='Разделитель главы:')
label.pack(padx=6, pady=6)
divider = tk.Entry(root)
divider.pack(padx=6, pady=3)

submit_button = tk.Button(root, text='Применить', command=submit)
submit_button.pack(padx=6, pady=10)
ready = tk.Label(root, wraplength=350)
ready.pack(padx=6, pady=6)

root.mainloop()
